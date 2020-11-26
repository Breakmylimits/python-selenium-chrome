from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from bs4 import BeautifulSoup as soup
import time
from docx import Document




def openpage(id):
    
    hight = 0
    
    url= 'https://twitter.com/'+id
    driver = webdriver.Chrome()
    driver.get(url)
    
    time.sleep(5)

    document = Document()

    document.add_heading('POST TWITTER', 0)

    p = document.add_paragraph(' ')

    for i in range(10):  
        page_html = driver.page_source
        
        data = soup(page_html, 'html.parser')
    
        tweet = data.findAll('div',{'class': 'css-901oao r-18jsvk2 r-1qd0xha r-a023e6 r-16dba41 r-ad9z0x r-bcqeeo r-bnwqim r-qvutc0'})
        user = data.findAll('div',{'class':'css-901oao css-bfa6kz r-m0bqgq r-18u37iz r-1qd0xha r-a023e6 r-16dba41 r-ad9z0x r-bcqeeo r-qvutc0'})
    
        result = []   

        driver.execute_script("window.scrollTo(0,{})".format(hight))
        hight += 20000
        time.sleep(3) 

        for i,tw in enumerate(zip(user,tweet)):
            i = i+1
            print(i)
            user =tw[0].text             
            print('user : '+user)
                
            comment = tw[1].text
            print('comment : '+comment)
            print('--------------------------------------------------')
            
            p.add_run('ชื่อผู้ใช้ : '+user+'\n').bold = True
            p.add_run('comment : '+comment+'\n\n')
            document.save('data Twitter.docx')

        
            
        
    driver.close()

    return     
     

    
#id =input('watch?v=fgP-2XJC0fc')



comment1 = openpage('search?q=%23แจงรายจ่ายม็อบด้วยจ้า&src=trend_click&vertical=trends')
print('don!')







#elem_search = driver.find_element_by_name('q')
#elem_search.clear()
#elem_search.send_keys('emi fukada')
#elem_search.send_keys(Keys.ENTER)
#css-901oao css-16my406 r-1qd0xha r-ad9z0x r-bcqeeo r-qvutc0
