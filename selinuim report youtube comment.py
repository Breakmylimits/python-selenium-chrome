from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from bs4 import BeautifulSoup as soup
import time
from docx import Document

def openpage(id):
    
    hight = 700
    driver = webdriver.Chrome()
    url= 'https://www.youtube.com/'+id
    driver.get(url)
    
    time.sleep(5)
    
    for i in range(5):     

        driver.execute_script("window.scrollTo(0,{})".format(hight))
        hight += 5000
        time.sleep(3)    
     
    page_html = driver.page_source
 
    data = soup(page_html, 'html.parser')
    
    tweet = data.findAll('div',{'class': 'style-scope ytd-expander'})
    user = data.findAll('a',{'class':'yt-simple-endpoint style-scope ytd-comment-renderer'})
    
    result = []
    
    
  
    for i,tw in enumerate(zip(user,tweet)):
        
        i = i+1
        print(i)
        user =tw[0].text
        user =user.replace('\n','')
        user =user.replace('             ','')
        #print([tw[0].text])
        print('user : '+user)
        print('https://www.youtube.com/'+tw[0]['href'])
        url = 'https://www.youtube.com/'+tw[0]['href']
        comment = tw[1].text
        comment =comment.replace('\n','')
        print('comment : '+comment)
        
        result.append({'user':user,'url':url,'comment':comment})


    driver.close()

    return result
#id =input('watch?v=fgP-2XJC0fc')



comment1 = openpage('watch?v=eEpcbMKgW0I')

document = Document()

document.add_heading('รายชื่อคอมเม้นต์', 0)

p = document.add_paragraph(' ')
for n in comment1:
    p.add_run('ชื่อผู้ใช้ : '+n['user']+'\n').bold = True
    p.add_run('url : '+n['url']+'\n')
    p.add_run('comment : '+n['comment']+'\n\n')


document.save('data.docx')
print('don!')

#elem_search = driver.find_element_by_name('q')
#elem_search.clear()
#elem_search.send_keys('emi fukada')
#elem_search.send_keys(Keys.ENTER)
#css-901oao css-16my406 r-1qd0xha r-ad9z0x r-bcqeeo r-qvutc0
